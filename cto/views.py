#from json.decoder import JSONDecoder
#from typing import Any, Dict
from django.conf.urls import url
from django.urls.conf import path
import docx
import locale
import json
import requests
import urllib
from urllib import request, parse


import io
from operator import index
from pathlib import Path
import os


# Para utilizar algunas de las funciones de la librería
from docx import Document
from docx.shared import Inches, Pt, Cm

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_BREAK_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn


from django.shortcuts import render, redirect, get_list_or_404
from django.views import generic
from django.views.generic import TemplateView, ListView, CreateView
from django.contrib.messages.views import SuccessMessageMixin

from django.urls import reverse_lazy
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.models import User
from django.http import HttpResponse, HttpResponseRedirect, HttpRequest, JsonResponse, HttpResponseServerError

from datetime import datetime, date
from django.contrib import messages
from django.db.models import Q
from django.contrib.auth import authenticate
from requests.api import get
from bases.views import SinPrivilegios
from bases.views import *
from .models import Departamento, Partes, Contratos, Doctos, Tipocontrato, Requisitos, Valida, Secuencia, Regimen, Ciclos, Puestos
from .forms import DepartamentoForm, PartesForm, ContratosForm, PuestosForm
from bases.views import SinPrivilegios
from django.core.files.storage import FileSystemStorage, get_storage_class
from api.serializer import TipocontratoSerializer


class VistaBaseCreate(SuccessMessageMixin, SinPrivilegios, generic.CreateView):
    context_object_name = 'obj'
    success_message = "Registro Agregado Satisfactoriamente"

    def form_valid(self, form):
        form.instance.uc = self.request.user
        return super().form_valid(form)


class VistaBaseEdit(SuccessMessageMixin, SinPrivilegios, generic.UpdateView):
    context_object_name = 'obj'
    success_message = "Registro Actualizado Satisfactoriamente"

    def form_valid(self, form):
        #form.instance.um = self.request.user.id
        return super().form_valid(form)


class DepartamentoView(SinPrivilegios, generic.ListView):
    model = Departamento
    template_name = "cto/departamento_list.html"
    context_object_name = "obj"
    permission_required = "cto.view_departamento"


class DepartamentoNew(VistaBaseCreate):
    model = Departamento
    template_name = "cto/departamento_form.html"
    form_class = DepartamentoForm
    success_url = reverse_lazy("cto:departamento_list")
    permission_required = "cto.add_departamento"


class DepartamentoEdit(VistaBaseEdit):
    model = Departamento
    template_name = "cto/departamento_form.html"
    form_class = DepartamentoForm
    success_url = reverse_lazy("cto:departamento_list")
    permission_required = "cto.change_departamento"


@login_required(login_url="/login/")
@permission_required("cto.change_departamento", login_url="/login/")
def departamentoInactivar(request, id):
    departamento = Departamento.objects.filter(pk=id).first()

    if request.method == "POST":
        if departamento:
            departamento.estado = not departamento.estado
            departamento.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


class PartesView(SinPrivilegios, generic.ListView):
    model = Partes
    template_name = "cto/partes_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:partes_list")
    permission_required = "cto.view_partes"

    def get_queryset(self):
        current_userx = self.request.user.id
        # print(current_userx)
        queryset = Partes.objects.filter(user_id=current_userx)
        xdepa = 2
        for part in queryset:
            xdepa = part.claveDepartamento_id

        # print(xdepa)

        querydep = Departamento.objects.all()
        # print(querydep)
        for depa in querydep:

            if depa.claveDepartamento == xdepa:
                # print(depa.claveDepartamento)
                xr1 = "0"+str(depa.rango1)
                xr2 = "0"+str(depa.rango2)
                # print(xdepa)
                # print(xr1)
                # print(xr2)
                return Partes.objects.filter(Q(claveDepartamento__gte=xr1),  Q(claveDepartamento__lte=xr2))
            # else:
            #    messages.ERROR (request, '¡Rango de Departamentos no registrado!')
            #    return HttpResponseRedirect('/admn/')

    def get_context_data(self, **kwargs):
        # Call the base implementation first to get a context
        context = super(PartesView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        return context


class PartesNew(VistaBaseCreate):
    model = Partes
    template_name = "cto/partes_form.html"
    form_class = PartesForm
    success_url = reverse_lazy("cto:partes_list")
    permission_required = "cto.add_partes"


class PartesEdit(VistaBaseEdit):
    model = Partes
    template_name = "cto/partes_form.html"
    form_class = PartesForm
    success_url = reverse_lazy("cto:partes_list")
    permission_required = "cto.change_partes"


@login_required(login_url="/login/")
@permission_required("cto.change_partes", login_url="/login/")
def partesInactivar(request, id):
    partes = Partes.objects.filter(pk=id).first()

    if request.method == "POST":
        if partes:
            partes.estado = not partes.estado
            partes.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


class ContratosView(SinPrivilegios, generic.ListView):
    model = Contratos
    template_name = "cto/contrato_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:contrato_list")
    permission_required = "cto.view_contratos"

    def get_queryset(self):
        xtipo = 0
        tipocontrato = Tipocontrato.objects.filter(marcatipoContrato=True)
        for x in tipocontrato:
            if x.marcatipoContrato == True:
                xtipo = x.id
                # print(xtipo)
        # r = requests.get('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print (r.content)
        # print (r.status_code)
        # print (r.headers)
        # print (r.json)
        # x=urllib.request.urlretrieve('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print(x)8

        current_userx = self.request.user.id
        #conditions = dict(current_user=current_userx, uc_id=self.request.user)
        #queryset = queryset.filter(**conditions)
        # return Contratos.objects.all()

        return Contratos.objects.filter(
            (Q(current_user=current_userx) | Q(uc_id=self.request.user)), Q(
                tipocontrato_id=xtipo)
        )
        # return SpyorEnc.objects.filter(
        #    Q(current_user=current_userx) | Q(uc_id=self.request.user) | Q(el_jefe=current_userx)
        # )

    def get_context_data(self, **kwargs):

        # Call the base implementation first to get a context
        context = super(ContratosView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        context['some_data'] = Partes.objects.all()
        context['some_data2'] = Departamento.objects.all()
        context['some_data3'] = Tipocontrato.objects.filter(
            marcatipoContrato=True)
        return context


class ContratosView2(SinPrivilegios, generic.ListView):
    model = Contratos
    template_name = "cto/contrato_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:contrato_list")
    permission_required = "cto.view_contratos"

    def get_queryset(self):
        xtipo = 0
        tipocontrato = Tipocontrato.objects.filter(marcatipoContrato=True)
        for x in tipocontrato:
            if x.marcatipoContrato == True:
                xtipo = x.id
                # print(xtipo)
        # r = requests.get('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print (r.content)
        # print (r.status_code)
        # print (r.headers)
        # print (r.json)
        # x=urllib.request.urlretrieve('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print(x)8

        current_userx = self.request.user.id
        #conditions = dict(current_user=current_userx, uc_id=self.request.user)
        #queryset = queryset.filter(**conditions)
        # return Contratos.objects.all()

        return Contratos.objects.filter(
            (Q(current_user=current_userx) | Q(uc_id=self.request.user)), Q(
                tipocontrato_id=xtipo)
        )
        # return SpyorEnc.objects.filter(
        #    Q(current_user=current_userx) | Q(uc_id=self.request.user) | Q(el_jefe=current_userx)
        # )

    def get_context_data(self, **kwargs):

        # Call the base implementation first to get a context
        context = super(ContratosView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        context['some_data'] = Partes.objects.all()
        context['some_data2'] = Departamento.objects.all()
        context['some_data3'] = Tipocontrato.objects.filter(
            marcatipoContrato=True)
        return context


@login_required(login_url='/login/')
@permission_required('cto.add_contratos', login_url='bases:sin_privilegios')
def contratos2(request, contrato_id=None):

    template_name = 'cto/contrato.html'
    detalle = {}
    #secuencia_data = {}
    #xUsuario = (request.user.id)
    # print(xUsuario)
    #print (contrato_id)

    if contrato_id:
        contrato = Contratos.objects.filter(estado=True, id=contrato_id)
        c = contrato.first()
        xUsuario = (c.parte2_id)
        # print(xUsuario)
        partes = Partes.objects.filter(estado=True, id=xUsuario)
        p = partes.first()
        #print (p)
        dx = p.claveDepartamento_id
    else:
        xUsuario = (request.user.username)
        xUsuario2 = (request.user.id)
        # print(xUsuario2)
        partes = Partes.objects.filter(estado=True, usuario=xUsuario)
        p = partes.first()
        #print (p)
        dx = p.claveDepartamento_id

    #d1 = p.claveDepartamento

    #rfcparte = (p.rfc)
    # print(rfcparte)

    # print(dx)

    requisitos = Requisitos.objects.filter(estado=True)
    departamentos = Departamento.objects.filter(
        estado=True, claveDepartamento=dx)
    departamentos2 = Departamento.objects.filter(estado=True)
    d2 = departamentos.first()
    # print(d2)
    d3 = (d2.claveDepartamento)
    # print(d3)
    #r1 = (d2.rango1)
    #r2 = (d2.rango2)
    r1 = "0"+str(d2.rango1)
    r2 = "0"+str(d2.rango2)
    secuencia = (d2.f001)

    # print(d2.testigoUsual1)
    # print(d2.testigoUsual2)
    #print (d3)
    # print(secuencia)
    #print (r1)
    #print (r2)

    partes3 = Partes.objects.filter(Q(estado=True),  Q(claveDepartamento_id__gte=r1), Q(
        claveDepartamento_id__lte=r2)).order_by('nombreParte')

    #print (partes3)

    partes2 = Partes.objects.filter(estado=True).order_by('nombreParte')
    #print (partes2)

    if secuencia:

        r1 = int(secuencia[:5])
        f1 = secuencia[5:8]

        r2 = int(secuencia[8:13])
        f2 = secuencia[13:16]

        r3 = int(secuencia[16:21])
        f3 = secuencia[21:24]

        if f1 == 'DIC':
            funcionario = Partes.objects.filter(user_id=r1)
        else:
            if f2 == 'DIC':
                funcionario = Partes.objects.filter(user_id=r2)
            else:
                if f3 == 'DIC':
                    funcionario = Partes.objects.filter(user_id=r3)
        # print(funcionario)

        a2 = funcionario.first()
        # print(a2)
        a3 = (a2.id)
        #print (a3)
        fun = Partes.objects.get(pk=a3)

        r4 = int(secuencia[24:29])
        f4 = secuencia[29:32]

        r5 = int(secuencia[32:37])
        f5 = secuencia[37:40]

        r6 = int(secuencia[40:45])
        f6 = secuencia[45:48]

        # print(r1)
        # print(f1)
        # print(r2)
        # print(f2)
        # print(r3)
        # print(f3)
        # print(r4)
        # print(f4)
        # print(r5)
        # print(f5)
        # print(r6)
        # print(f6)
    else:
        #print ("secuencia no asignada")
        messages.error(request, '¡Secuencia administrativa no asignada!')
        #messages.error = ('secuencia no asignada')
        return HttpResponseRedirect('/cto/contratos/')

    if request.method == "GET":
        enc = Contratos.objects.filter(pk=contrato_id).first()
        tipocontratox = Tipocontrato.objects.filter(marcatipoContrato=True)
        tipocontratox = tipocontratox.first()

        tipocontrato = tipocontratox.id
        if not enc:
            encabezado = {
                'id': "",
                'uc_id': "",
                'tipocontrato': tipocontrato,
                'datecontrato': datetime.today(),
                'datecontrato_ini': "",
                'datecontrato_fin': "",
                'parte1': 164,
                'enCalidadDe1': '"CLIENTE"',
                'parte2': "",
                'enCalidadDe2': '"PRESTADOR DE SERVICIOS"',
                'lugarContrato': "",
                'ciudadContrato': "Mérida",
                'estadoContrato': "Yucatán",
                'paisContrato': "México",
                'importeContrato': 0.00,
                'npContrato': 1,
                'imppContrato': 0.00,
                'vhppContrato': 285.00,
                'totalhorasContrato': "",
                'testigoContrato1': d2.testigoUsual1,
                'testigoContrato2': d2.testigoUsual2,
                'versionContrato': "",

                'status': "CAP",
                'rcap': xUsuario2,


                'rstep1': r1,
                'rstep2': r2,
                'rstep3': r3,
                'rstep4': r4,
                'rstep5': r5,
                'rstep6': r6,

                'astep1': f1,
                'astep2': f2,
                'astep3': f3,
                'astep4': f4,
                'astep5': f5,
                'astep6': f6,
                'devuelto_por': "",

            }
            detalle = None
        else:

            encabezado = {
                'id': enc.id,
                'uc_id': enc.uc_id,
                'tipocontrato': enc.tipocontrato,
                'datecontrato': enc.datecontrato,
                'datecontrato_ini': enc.datecontrato_ini,
                'datecontrato_fin': enc.datecontrato_fin,
                'parte1': enc.parte1,
                'enCalidadDe1': enc.enCalidadDe1,
                'parte2': enc.parte2,
                'enCalidadDe2': enc.enCalidadDe2,
                'lugarContrato': enc.lugarContrato,
                'ciudadContrato': enc.ciudadContrato,
                'estadoContrato': enc.estadoContrato,
                'paisContrato': enc.paisContrato,
                'importeContrato': enc.importeContrato,
                'npContrato': enc.npContrato,
                'imppContrato': enc.imppContrato,
                'vhppContrato': enc.vhppContrato,
                'totalhorasContrato': enc.totalhorasContrato,
                'testigoContrato1': enc.testigoContrato1,
                'testigoContrato2': enc.testigoContrato2,
                'versionContrato': enc.versionContrato,

                'status': enc.status,
                'rcap': enc.rcap,


                'rstep1': enc.rstep1,
                'rstep2': enc.rstep2,
                'rstep3': enc.rstep3,
                'rstep4': enc.rstep4,
                'rstep5': enc.rstep5,
                'rstep6': enc.rstep6,

                'astep1': enc.astep1,
                'astep2': enc.astep2,
                'astep3': enc.astep3,
                'astep4': enc.astep4,
                'astep5': enc.astep5,
                'astep6': enc.astep6,

                'fstep1': enc.fstep1,
                'fstep2': enc.fstep2,
                'fstep3': enc.fstep3,
                'fstep4': enc.fstep4,
                'fstep5': enc.fstep5,
                'fstep6': enc.fstep6,

                'cstep1': enc.cstep1,
                'cstep2': enc.cstep2,
                'cstep3': enc.cstep3,
                'cstep4': enc.cstep4,
                'cstep5': enc.cstep5,
                'cstep6': enc.cstep6,

                'devuelto_por': enc.devuelto_por,
                'current_user': enc.current_user,

            }

        detalle = Doctos.objects.filter(contrato=enc)

        contexto = {"requi": requisitos, "fun": funcionario, "fun2": partes2, "fun3": partes3, "enc": encabezado, "det": detalle,
                    "departamentos": departamentos, "funcionarios": partes, "departamentos2": departamentos2, "tipocont": tipocontratox, }

        return render(request, template_name, contexto)

    if request.method == "POST":

        tipocontratox = Tipocontrato.objects.filter(marcatipoContrato=True)
        tipocontratox = tipocontratox.first()

        tipocontrato = tipocontratox.id
        # print(tipocontrato)
        datecontrato = request.POST.get("datecontrato")
        datecontrato_ini = request.POST.get("enc_datecontrato_ini")
        datecontrato_fin = request.POST.get("enc_datecontrato_fin")

        parte1 = 164
        parte2 = request.POST.get("enc_nombreParte")

        enCalidadDe1 = tipocontratox.enCalidadDe1
        enCalidadDe2 = tipocontratox.enCalidadDe2

        lugarContrato = request.POST.get("lugarContrato")
        ciudadContrato = "Mérida"
        estadoContrato = "Yucatán"
        paisContrato = "México"
        importeContrato = request.POST.get("enc_importeContrato")
        npContrato = request.POST.get("enc_npContrato")
        imppContrato = request.POST.get("enc_imppContrato")

        vhppContrato = 285
        totalhorasContrato = request.POST.get("enc_totalhorasContrato")
        testigoContrato1 = d2.testigoUsual1
        testigoContrato2 = d2.testigoUsual2
        versionContrato = request.POST.get("versionContrato")
        status = request.POST.get("status")

        fun = Partes.objects.get(pk=a3)
        tip = Tipocontrato.objects.get(pk=tipocontrato)

        if parte2:
            suj = Partes.objects.get(pk=parte2)

        if not contrato_id:
            enc = Contratos(
                tipocontrato=tip,
                datecontrato=datecontrato,
                datecontrato_ini=datecontrato_ini,
                datecontrato_fin=datecontrato_fin,

                parte1=parte1,
                parte2=suj,
                ciudadContrato=ciudadContrato,
                estadoContrato=estadoContrato,
                paisContrato=paisContrato,

                enCalidadDe1=enCalidadDe1,
                enCalidadDe2=enCalidadDe2,

                importeContrato=importeContrato,
                npContrato=npContrato,
                imppContrato=imppContrato,
                vhppContrato=vhppContrato,
                totalhorasContrato=totalhorasContrato,
                testigoContrato1=d2.testigoUsual1,
                testigoContrato2=d2.testigoUsual2,



                status=status,
                rcap=xUsuario2,
                current_user=xUsuario2,
                rstep1=r1,
                rstep2=r2,
                rstep3=r3,
                rstep4=r4,
                rstep5=r5,
                rstep6=r6,

                astep1=f1,
                astep2=f2,
                astep3=f3,
                astep4=f4,
                astep5=f5,
                astep6=f6,
                devuelto_por=False,
            )
            if enc:
                enc.save()
                contrato_id = enc.id
        else:
            enc = Contratos.objects.filter(pk=contrato_id).first()
            if enc:

                enc.funcionario = fun

                enc.save()

        if not id:
            messages.error(
                request, 'No Puedo Continuar No Pude Detectar No. de Contrato')
            return redirect("cto:contrato_list")

        documento = request.POST.get("documento")
        comentarioDocto = request.POST.get("comentarioDocto")
        type(documento)
        req = Requisitos.objects.get(pk=documento)
        vigenciaFinDocto = request.POST.get("enc_vigenciaFinDocto")

        pdf = request.POST.get('pdf2')

        uploaded_file = request.FILES['pdf']
        # print(uploaded_file)

        fs = FileSystemStorage()
        # print(fs)
        name = fs.save(uploaded_file.name, uploaded_file)
        # print(name)
        pdf = name

        if vigenciaFinDocto != "":

            det = Doctos(
                contrato=enc,
                documento=req,
                comentarioDocto=comentarioDocto,
                pdf=pdf,
                vigenciaFinDocto=vigenciaFinDocto,

            )

        if vigenciaFinDocto == "":

            det = Doctos(
                contrato=enc,
                documento=req,
                comentarioDocto=comentarioDocto,
                pdf=pdf,


            )

        if det:
            det.save()

            return redirect("cto:contrato_edit", contrato_id=contrato_id)

    return render(request, template_name, contexto)


class ContratosEdit(VistaBaseEdit):
    model = Contratos
    template_name = "cto/contrato_form.html"
    form_class = ContratosForm
    success_url = reverse_lazy('cto:contrato_list')
    permission_required = "cto.change_contratos"
    context_object_name = 'obj'


login_required(login_url="/login/")


@permission_required("cto.change_contratos", login_url="/login/")
def coverletter_export(request, id):

    contratos = Contratos.objects.filter(pk=id).first()  # Contrato en curso
    # Información del tipo de contrato
    tipoc = Tipocontrato.objects.get(id=contratos.tipocontrato_id)
    # print(tipoc.id)
    # Validacion de información completa
    valic = Valida.objects.filter(tipocontrato_id=tipoc.id)
    # Datos del contratado **sujeto del contrato
    partes = Partes.objects.get(id=contratos.parte2_id)
    patron = Partes.objects.get(id=164)  # Datos del contratante
    secue = Secuencia.objects.filter(
        tipocontrato_id=tipoc.id).first()  # Primer parrafo del contrato
    ciclo = Ciclos.objects.filter(
        ciclo_actual=True).first()  # Ciclo escolar actual
    # print(tipoc.id)
    # print(secue)
    # Régimen fiscal del contratado
    regimen = Regimen.objects.get(id=partes.regfiscalParte_id)
    replegal = Partes.objects.get(id=165)  # Datos del contratante
    letras = numero_to_letras(contratos.importeContrato.amount)
    
    if contratos.imppContrato:
        pagolet = numero_to_letras(contratos.imppContrato.amount)
        currency2 = "${:,.2f}".format(contratos.imppContrato.amount)
    else:
        pagolet = 0
        currency2 = ""
    currency = "${:,.2f}".format(contratos.importeContrato.amount)
    document = Document()
    locale.setlocale(locale.LC_TIME, "es-MX")
    puesto = Puestos.objects.filter(nombrePuesto=partes.clavePuesto).first()

    #footer_para.paragraph_format.page_break_before = True

    # set up font
    font = document.styles['Normal'].font
    font.name = 'Calibri'
    font.bold = True
    font.size = Pt(16)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style(
        'CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri'

    # set up margins
    sections = document.sections
    for section in sections:

        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(0)

    header = document.sections[0].header
    htable = header.add_table(1, 2, Inches(5))
    #htable.style = "TableGrid"

    for row in htable.rows:
        row.height = Inches(1.0)

    htable.alignment = WD_TABLE_ALIGNMENT.LEFT
    htab_cells = htable.rows[0].cells

    ht0 = htab_cells[0].paragraphs[0]
    paragraph_format = ht0.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.left_indent = Pt(0)

    kh = ht0.add_run()

    THIS_FOLDER = os.path.dirname(os.path.abspath(__file__))

    my_file = os.path.join(THIS_FOLDER, 'logo.png')

    kh.add_picture(my_file, width=Inches(1.00))
    ht0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ht1 = htab_cells[1].add_paragraph("ESCUELA MODELO, S.C.P.")
    htable.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    htable.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Get the user's fullname
    # if request.user.get_full_name():
    #   document_data_full_name = request.user.get_full_name()
    # else:
    #    document_data_full_name = "[NOMBRE] [APELLIDOS]"

    p001 = document.add_paragraph()
    p001.add_run(tipoc.tituloContrato, style='CommentsStyle')
    p001.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p002 = document.add_paragraph()

    textox = tipoc.textoinicialContrato

    if partes.tituloParte:
        xnombreParte = partes.tituloParte + " " + partes.nombreParte
    else:
        if partes.personaParte == 1:
            xnombreParte = "**********" + " " + partes.nombreParte
        else:
            xnombreParte = partes.nombreParte
        #messages.info(request, message="Registrar: Título del Sujeto del Contrato")

    if partes.curp:
        xcurp = partes.curp[10:11]
    else:
        xcurp = "X"

    if xcurp == "H":
        xelolaParte = "EL"
        xenCalidadDe2 = tipoc.enCalidadDe2
        yenCalidadDe2 = tipoc.enCalidadDe2

    else:
        xelolaParte = "LA"
        xenCalidadDe2 = tipoc.enCalidadDe2f
        yenCalidadDe2 = tipoc.enCalidadDe2f

    xenCalidadDe2 = textox.replace("@enCalidadDe2", xenCalidadDe2)

    textox = xenCalidadDe2
    xenCalidadDe1 = textox.replace("@enCalidadDe1", tipoc.enCalidadDe1)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@elolaParte", xelolaParte)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@nombreParte2", xnombreParte)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@nombreParte", patron.nombreParte)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@tituloParteRL", replegal.tituloParte)

    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace(
        "@idrep_legalParte", replegal.tituloParte + " " + replegal.nombreParte)

    p002.add_run(xenCalidadDe1, style='CommentsStyle').bold = False
    p002.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if tipoc.id == 1 or tipoc.id == 2 or tipoc.id == 3:
        p003 = document.add_paragraph()
        p003.add_run("CLÁUSULAS", style='CommentsStyle').bold = True
        p003.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if tipoc.id == 7 or tipoc.id == 8 or tipoc.id == 4 or tipoc.id == 5 or tipoc.id == 11 or tipoc.id == 13:
        p003 = document.add_paragraph()
        p003.add_run("DECLARACIONES", style='CommentsStyle').bold = True
        p003.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # print(secue.id)
    #secue = Secuencia.objects.get(id=1)
    p004 = document.add_paragraph()
    textosecuex = secue.identificador + ".- " + secue.textoSecuencia
    textosecuex = textosecuex.replace("@enCalidadDe2", yenCalidadDe2)
    # print(yenCalidadDe2)
    textosecuex = textosecuex.replace("@enCalidadDe1", tipoc.enCalidadDe1)
    # print(textosecuex)
    p004.add_run(textosecuex, style='CommentsStyle').bold = True
    paragraph_format = p004.paragraph_format
    paragraph_format.left_indent = Inches(0.0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if tipoc.id == 1:
        nums = (2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19,
                20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36)

    if tipoc.id == 7:
        nums = (314, 315, 316, 317, 318, 319, 320, 321, 322, 323, 324, 325, 326, 327, 328, 338, 339, 340, 341, 342, 343, 344, 345, 346,
                347, 348, 349, 350, 352, 353, 354, 355, 356, 357, 358, 359, 360, 361, 362, 363, 364, 365, 366, 367, 368, 369, 370, 371, 372, 373, 374, 375, 376, 377, 378, 379, 380, 381)

    if tipoc.id == 8:
        nums = (384, 385, 386, 387, 388, 389, 390, 391, 392, 393, 394, 395, 396, 397, 398, 399, 400, 401, 402, 403, 404, 405, 406, 407, 408, 409, 410, 411, 412, 413, 414, 415, 416, 417, 418, 419,
                420, 421, 422, 423, 424, 425, 426, 427, 428, 429, 430, 431, 432, 433, 434, 435, 436, 437, 438, 439, 440, 441, 442, 443, 444, 445, 446, 447, 448, 449, 450, 451, 452, 453, 454, 455, 456, 457)

    if tipoc.id == 9:
        nums = (460, 461, 462, 463, 464, 465, 466, 467, 468, 469, 470, 471, 472, 473, 474, 475, 476, 477, 478, 479, 480, 481, 482, 483, 484, 485, 486, 487, 488, 489, 490, 491, 492, 493, 494, 495,
                496, 497, 498, 499, 500, 501, 502, 503, 504, 505, 506, 507, 508, 509, 510, 511, 512, 513, 514, 515, 516, 517, 518, 519, 520, 521, 522, 523, 524, 525, 526, 527, 528, 529, 530, 531, 532, 533)

    if tipoc.id == 10:
        nums = (535, 536, 537, 538, 539, 540, 541, 542, 543, 544, 545, 546, 547, 548, 549, 550, 551, 552, 553, 554, 555, 556, 557, 558, 559, 560, 561, 562, 563, 564, 565, 566,
                567, 568, 569, 570, 571, 572, 573, 574, 575, 576, 577, 578, 579, 580, 581, 582, 583, 584, 585, 586, 587, 588, 589, 590, 591, 592, 593, 594, 595, 596, 597, 598, 599, 600)
    #secue = Secuencia.objects.get(id=2)

    if tipoc.id == 4:
        nums = (129, 130, 131, 132, 133, 134, 135, 136, 137, 138, 139, 140, 141, 142, 143, 144, 145, 146, 147, 148, 149, 150, 151, 152, 153, 154, 155, 156, 157,
                158, 159, 160, 161, 162, 163, 164, 165, 166, 167, 168, 169, 170, 171, 172, 173, 174, 175, 176, 177, 178, 179, 180, 181, 182, 183, 184, 185, 186, 187)

    if tipoc.id == 5:
        nums = (190, 191, 192, 193, 194, 195, 196, 197, 198, 199, 200, 201, 202, 203, 204, 205, 206, 207, 208, 209, 210, 211, 212, 213, 214, 215, 216, 217,
                218, 219, 220, 221, 222, 223, 224, 225, 226, 227, 228, 229, 230, 231, 232, 233, 234, 235, 236, 237, 238, 239, 240, 241, 242, 243, 244, 245, 246)

    if tipoc.id == 6:
        nums = (249, 250, 251, 252, 253, 254, 255, 256, 257, 258, 259, 260, 261, 262, 263, 264, 265, 266, 267, 268, 269, 270, 271, 272, 273, 274, 275, 276, 277, 278, 279,
                280, 281, 282, 283, 284, 285, 286, 287, 288, 289, 290, 291, 292, 293, 294, 295, 296, 297, 298, 299, 300, 301, 302, 303, 304, 305, 306, 307, 308, 309, 310, 311, 312)

    if tipoc.id == 2:
        nums = (38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57,
                58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79)

    if tipoc.id == 3:
        nums = (82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100, 101, 102, 103, 104, 105,
                106, 107, 108, 109, 110, 111, 112, 113, 114, 115, 116, 117, 118, 119, 120, 121, 122, 123, 124, 125, 126)

    if tipoc.id == 11:
        nums = (604, 605, 606, 607, 608, 609, 610, 611, 612, 613, 614, 615, 616, 617, 618, 619, 620, 621, 622, 623, 624, 625,
                626, 627, 628, 629,  630, 631, 632, 633, 634, 635, 636, 637, 638, 639, 640, 641, 642, 643, 644, 645, 646, 647, 648)

    if tipoc.id == 13:
        nums = (649, 650)

    for n in nums:
        # print(n)
        secue = Secuencia.objects.get(id=n)
        if secue.nivel2 == 0:

            secue = Secuencia.objects.get(id=n)
            p005 = document.add_paragraph()
            if secue.identificador != "," and secue.identificador != ".":
                textosecue = secue.identificador + ".- " + secue.textoSecuencia
            else:
                textosecue = secue.textoSecuencia

            textosecue = textosecue.replace(
                "@enCalidadDe1", tipoc.enCalidadDe1)

            if partes.personaParte == 1:
                xcurp = partes.curp[10:11]
                if xcurp == "H":
                    xelolaParte = "EL"
                    xenCalidadDe2 = tipoc.enCalidadDe2
                    textosecue = textosecue.replace(
                        "@enCalidadDe2", xenCalidadDe2)
                else:
                    xelolaParte = "LA"
                    xenCalidadDe2 = tipoc.enCalidadDe2f
                    textosecue = textosecue.replace(
                        "@enCalidadDe2", xenCalidadDe2)

            else:
                xelolaParte = ""
                xenCalidadDe2 = tipoc.enCalidadDe2
                textosecue = textosecue.replace("@enCalidadDe2", xenCalidadDe2)

            if partes.personaParte == 1:
                if partes.fecha_ingreso:
                    # print(secue.id)
                    textosecue = textosecue.replace(
                        "@fechaingreso", partes.fecha_ingreso.strftime("%d de %B de %Y"))
                    
                
                else:
                    textosecue = textosecue.replace(
                        "@fechaingreso", "**********")
                    
               
                if puesto.funcionesPuesto:
                    # print(secue.id)
                    
                    textosecue = textosecue.replace(
                        "@actividadesPuesto", puesto.funcionesPuesto)
                   
                    
                else:
                    textosecue = textosecue.replace(
                        "@actividadesPuesto", "**********")
                    
                if partes.actividadesParte:
                    # print(secue.id)
                    textosecue = textosecue.replace(
                        "@actividadesContrato", partes.actividadesParte)
                else:
                    textosecue = textosecue.replace("@actividadesContrato", "")

                if partes.titulo_profParte:
                    textosecue = textosecue.replace(
                        "@titulo_profParte", partes.titulo_profParte)
                else:
                    textosecue = textosecue.replace(
                        "@titulo_profParte", "**********")

                if partes.clavePuesto:
                    # print(secue.id)
                    textosecue = textosecue.replace(
                        "@clavePuesto", puesto.nombrePuesto)
                else:
                    textosecue = textosecue.replace(
                        "@clavePuesto", "**********")

            textosecue = textosecue.replace(
                "@totalhorasContrato", str(contratos.totalhorasContrato))
            textosecue = textosecue.replace("@importeContrato", currency)
            textosecue = textosecue.replace("@letras", "(" + letras + ")")

            textosecue = textosecue.replace(
                "@datecontrato_ini", contratos.datecontrato_ini.strftime("%d de %B de %Y"))
            if contratos.datecontrato_fin:
                textosecue = textosecue.replace(
                    "@datecontrato_fin", contratos.datecontrato_fin.strftime("%d de %B de %Y"))
            textosecue = textosecue.replace(
                "@datecontratox", contratos.datecontrato.strftime("%d de %B de %Y"))
            if xcurp == "M":
                textosecue = textosecue.replace(
                    "CATEDRÁTICO UNIVERSITARIO Y ASESOR", "CATEDRÁTICA UNIVERSITARIA Y ASESORA")
                textosecue = textosecue.replace("este último", "esta última")
            if secue.identificador == ".":
                p005.add_run(textosecue, style='CommentsStyle').bold = False
            else:
                p005.add_run(textosecue, style='CommentsStyle').bold = True
            paragraph_format = p005.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.left_indent = Inches(0.0)
            print(len(textosecue))
            print(textosecue)
           
            if len(textosecue) > 100 and len(textosecue) <1500 :
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
              
            else:
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
               

        else:

            secue = Secuencia.objects.get(id=n)
            p006 = document.add_paragraph()
            textosecue = secue.identificador + ".- " + secue.textoSecuencia

            if partes.datos_actaconstParte:
                textosecue = textosecue.replace(
                    "@datos_actaconstParte", partes.datos_actaconstParte)

            if partes.curp:
                textosecue = textosecue.replace("@curp", partes.curp)

            if partes.titulo_profParte:
                textosecue = textosecue.replace(
                    "@titulo_profParte", partes.titulo_profParte)
            else:
                textosecue = textosecue.replace(
                    "@titulo_profParte", "**********")

            textosecue = textosecue.replace(
                "@enCalidadDe1", tipoc.enCalidadDe1)
            textosecue = textosecue.replace("@enCalidadDe2", xenCalidadDe2)

            if partes.universidadParte:
                textosecue = textosecue.replace(
                    "@universidadParte", partes.universidadParte)
            else:
                textosecue = textosecue.replace(
                    "@universidadParte", "**********")

            if partes.cedula_profParte:
                textosecue = textosecue.replace(
                    "@cedula_profParte", partes.cedula_profParte)
            else:
                textosecue = textosecue.replace(
                    "@cedula_profParte", "**********")

            if ciclo.ciclo_actual:
                textosecue = textosecue.replace(
                    "@CicloParte", ciclo.descripcionCiclo)
            else:
                textosecue = textosecue.replace("@CicloParte", "**********")

            if partes.rfc:
                textosecue = textosecue.replace("@rfc", partes.rfc)
            else:
                textosecue = textosecue.replace("@rfc", "**********")

            if regimen.nombreRegimen:
                textosecue = textosecue.replace(
                    "@regfiscalParte", regimen.nombreRegimen)
            else:
                textosecue = textosecue.replace(
                    "@regfiscalParte", "**********")

            if partes.domicilioParte:
                textosecue = textosecue.replace(
                    "@domicilioParte", partes.domicilioParte)
            else:
                textosecue = textosecue.replace(
                    "@domicilioParte", "**********")

            if partes.nacionalidadParte:
                textosecue = textosecue.replace(
                    "@nacionalidadParte", partes.nacionalidadParte)
            else:
                textosecue = textosecue.replace(
                    "@nacionalidadParte", "**********")

            ano = partes.rfc[4:6]
            mes = partes.rfc[6:8]
            dia = partes.rfc[8:10]
            xano = int(ano) + 1900

            if partes.personaParte == 1:
                fecha_nacimiento = date(xano, int(mes), int(dia))
                edad = calcular_edad_anos(fecha_nacimiento)
                #print(f'la edad es {edad} años')

                if edad:
                    textosecue = textosecue.replace(
                        "@edadParte", " "+str(edad))
                else:
                    textosecue = textosecue.replace("@edadParte", "**********")

                sexo = partes.curp[10:11]

                if sexo == "H":
                    textosecue = textosecue.replace("@sexoParte", "MASCULINO")
                else:
                    if sexo == "M":
                        textosecue = textosecue.replace(
                            "@sexoParte", "FEMENINO")
                    else:
                        textosecue = textosecue.replace(
                            "@sexoParte", "**********")

                if partes.estadocivilParte:
                    textosecue = textosecue.replace(
                        "@estadocivilParte", partes.estadocivilParte)
                else:
                    textosecue = textosecue.replace(
                        "@estadocivilParte", "**********")

            textosecue = textosecue.replace(
                "@datecontrato_ini", contratos.datecontrato_ini.strftime("%d de %B de %Y"))
            if contratos.datecontrato_fin:
                textosecue = textosecue.replace(
                    "@datecontrato_fin", contratos.datecontrato_fin.strftime("%d de %B de %Y"))
            textosecue = textosecue.replace(
                "@datecontratox", contratos.datecontrato.strftime("%d de %B de %Y"))

            textosecue = textosecue.replace(
                "@idrep_legalParte", replegal.tituloParte + " " + replegal.nombreParte)
            textosecue = textosecue.replace("@RPImssParteC", "8401667310-9 ")
            textosecue = textosecue.replace(
                "@enCalidadDe1", contratos.enCalidadDe1)
            textosecue = textosecue.replace(
                "@enCalidadDe2", contratos.enCalidadDe2)
            textosecue = textosecue.replace(
                "@domicilioPatron", patron.domicilioParte)
            textosecue = textosecue.replace("@importeContrato", currency)
            textosecue = textosecue.replace("@letras", "(" + letras + ")")
            textosecue = textosecue.replace(
                "@npContrato", str(contratos.npContrato))

            textosecue = textosecue.replace(
                "@totalhorasContrato", str(contratos.totalhorasContrato))
            if currency2:
                textosecue = textosecue.replace("@imppContrato", currency2)
                textosecue = textosecue.replace(
                    "@pagolet", "(" + pagolet + ")")
            if xcurp == "M":
                textosecue = textosecue.replace("mexicano", "mexicana")
                textosecue = textosecue.replace(
                    "un profesionista", "una profesionista")
                textosecue = textosecue.replace("inscrito", "inscrita")

            if puesto.funcionesPuesto:
                # print(secue.id)
                textosecue = textosecue.replace(
                    "@actividadesPuesto", puesto.funcionesPuesto)
              
                textosecue = textosecue.replace(
                    "@actividadesPuesto", "**********")

            p006.add_run(textosecue, style='CommentsStyle').bold = False
            paragraph_format = p006.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(3)
            paragraph_format.left_indent = Inches(0.4)
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if partes.personaParte == 1:
                xcurp = partes.curp[10:11]

    dtable = document.add_table(rows=3, cols=2)
    dtable.style = "TableNormal"
    dtable.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    dtable.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    dtab_cells = dtable.rows[0].cells
    dt1 = dtab_cells[0].text = ''
    dt1 = dtab_cells[0].paragraphs[0].add_run(
        tipoc.enCalidadDe1).font.size = Pt(11)
    dt1 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt1 = dtab_cells[1].text = ''
    dt1 = dtab_cells[1].paragraphs[0].add_run(xenCalidadDe2).font.size = Pt(11)
    dt1 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable.rows[1].cells
    dtable.rows[1].height_rule = WD_ROW_HEIGHT.EXACTLY
    dtable.rows[1].height = 342900
    dt1 = dtab_cells[0].text = ''
    dt1 = dtab_cells[0].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt1 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt1 = dtab_cells[1].text = ''
    dt1 = dtab_cells[1].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt1 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable.rows[2].cells
    dt2 = dtab_cells[0].text = ''
    dt2 = dtab_cells[0].paragraphs[0].add_run(
        replegal.tituloParte + " " + replegal.nombreParte).font.size = Pt(11)
    dt2 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt2 = dtab_cells[1].text = ''

    if partes.tituloParte:
        dt2 = dtab_cells[1].paragraphs[0].add_run(
            partes.tituloParte + " " + partes.nombreParte).font.size = Pt(11)
    else:
        if partes.personaParte == 1:
            dt2 = dtab_cells[1].paragraphs[0].add_run(
                "**********" + " " + partes.nombreParte).font.size = Pt(11)
        else:
            dt2 = dtab_cells[1].paragraphs[0].add_run(
                partes.nombreParte).font.size = Pt(11)

    dt2 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtable = document.add_table(rows=1, cols=2)

    dtable1 = document.add_table(rows=3, cols=2)
    dtable1.style = "TableNormal"
    dtable1.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable1.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable1.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable1.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable1.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    dtable1.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    dtab_cells = dtable1.rows[0].cells
    dt3 = dtab_cells[0].text = ''
    dt3 = dtab_cells[0].paragraphs[0].add_run("TESTIGO").font.size = Pt(11)
    dt3 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt3 = dtab_cells[1].text = ''
    dt3 = dtab_cells[1].paragraphs[0].add_run("TESTIGO").font.size = Pt(11)
    dt3 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable1.rows[1].cells
    dtable1.rows[1].height_rule = WD_ROW_HEIGHT.EXACTLY
    dtable1.rows[1].height = 342900
    dt4 = dtab_cells[0].text = ''
    dt4 = dtab_cells[0].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt4 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt4 = dtab_cells[1].text = ''
    dt4 = dtab_cells[1].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt4 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable1.rows[2].cells
    dt5 = dtab_cells[0].text = ''
    dt5 = dtab_cells[0].paragraphs[0].add_run(
        contratos.testigoContrato1).font.size = Pt(11)
    dt5 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt5 = dtab_cells[1].text = ''
    dt5 = dtab_cells[1].paragraphs[0].add_run(
        contratos.testigoContrato2).font.size = Pt(11)
    dt5 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    document.add_paragraph()
    xtable = document.add_table(rows=1, cols=2)
    xtable.style = "TableGrid"

    xtab_cells = xtable.rows[0].cells
    xt1 = xtab_cells[1].text = ''

    if partes.tituloParte:
        xt1 = xtab_cells[1].paragraphs[0].add_run(
            partes.tituloParte + " " + partes.nombreParte).font.size = Pt(8)
    else:
        if partes.personaParte == 1:
            xt1 = xtab_cells[1].paragraphs[0].add_run(
                "**********" + " " + partes.nombreParte).font.size = Pt(8)
        else:
            xt1 = xtab_cells[1].paragraphs[0].add_run(
                partes.nombreParte).font.size = Pt(8)

    xt1 = xtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    xt1 = xtab_cells[0].text = ''
    xt1 = xtab_cells[0].paragraphs[0].add_run(
        "DATOS PARA CONTROL EN EL ÁREA DE RECURSOS HUMANOS").font.size = Pt(8)
    xt1 = xtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytable = document.add_table(rows=2, cols=5)
    ytable.style = "TableGrid"

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[0].text = ''
    yt1 = ytab_cells[0].paragraphs[0].add_run(
        "Total Contrato en Horas").font.size = Pt(8)
    yt1 = ytab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[1].text = ''
    yt1 = ytab_cells[1].paragraphs[0].add_run(
        "Total Contrato en $").font.size = Pt(8)
    yt1 = ytab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[2].text = ''
    yt1 = ytab_cells[2].paragraphs[0].add_run(
        "Clave del Depto.").font.size = Pt(8)
    yt1 = ytab_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[3].text = ''
    yt1 = ytab_cells[3].paragraphs[0].add_run(
        "Ingreso o Reingreso").font.size = Pt(8)
    yt1 = ytab_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[4].text = ''
    yt1 = ytab_cells[4].paragraphs[0].add_run("Versión").font.size = Pt(8)
    yt1 = ytab_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[0].text = ''
    yt2 = ytab_cells[0].paragraphs[0].add_run(
        str(contratos.totalhorasContrato)).font.size = Pt(8)
    yt2 = ytab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[1].text = ''
    yt2 = ytab_cells[1].paragraphs[0].add_run(currency).font.size = Pt(8)
    yt2 = ytab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[2].text = ''
    yt2 = ytab_cells[2].paragraphs[0].add_run(
        str(partes.claveDepartamento)).font.size = Pt(8)
    yt2 = ytab_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[3].text = ''
    yt2 = ytab_cells[3].paragraphs[0].add_run(" ").font.size = Pt(8)
    yt2 = ytab_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[4].text = ''
    yt2 = ytab_cells[4].paragraphs[0].add_run(
        "EMODELO 17 v5 enero 2019").font.size = Pt(8)
    yt2 = ytab_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Added new section for assigning different footer on each page.
    new_section = document.sections[0].footer
    sectPr = new_section._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    # Get footer-area of the recent section in document
    new_footer = document.sections[0].footer
    new_footer.is_linked_to_previous = False

    footer_para = new_footer.paragraphs[0]
    run_footer = footer_para.add_run("Pág- ").font.size = Pt(11)
    run_footer = footer_para.add_run(pgNumType)
    _add_number_range(run_footer)
    font = run_footer.font
    font.name = 'Arial'
    font.size = Pt(11)
    run_footer = new_footer.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Print the user's name
    #document_elements_heading = document.add_heading(document_data_full_name, 0)
    #document_elements_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Print biography and careerpath

    # Add empty paragraph
    # document.add_paragraph()

    # Sincerely and name
    #document.add_paragraph("Atentamente,\n" + document_data_full_name)

    document_data = io.BytesIO()
    # print(document_data)
    document.save(document_data)
    document_data.seek(0)
    response = HttpResponse(
        document_data.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename = "Contrato.docx"'
    response["Content-Encoding"] = "UTF-8"
    return response


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def contratosAvanza(request, id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.method == "POST":

        if not contratos.fcap:
            contratos.fcap = datetime.today()
            contratos.current_user = contratos.rstep1
            contratos.status = contratos.astep1
            contratos.devuelto_por = contratos.rcap
            contratos.save()
            return HttpResponse("OK")
        # return HttpResponse("FAIL")
        if not contratos.fstep1:
            contratos.fstep1 = datetime.today()
            contratos.current_user = contratos.rstep2
            contratos.status = contratos.astep2
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep2:
            contratos.fstep2 = datetime.today()
            contratos.current_user = contratos.rstep3
            contratos.status = contratos.astep3
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep3:
            contratos.fstep3 = datetime.today()
            contratos.current_user = contratos.rstep4
            contratos.status = contratos.astep4
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep4:
            contratos.fstep4 = datetime.today()
            contratos.current_user = contratos.rstep5
            contratos.status = contratos.astep5
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep5:
            contratos.fstep5 = datetime.today()
            contratos.current_user = contratos.rstep6
            contratos.status = contratos.astep6
            contratos.save()
            return HttpResponse("OK")

        if not contratos.fstep6:
            contratos.fstep6 = datetime.today()
            #contratos.current_user = contratos.rstep6
            contratos.status = "FIN"
            contratos.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def contratosDevuelve(request, id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.method == "POST":
        if not contratos.fstep6:
            contratos.fcap = contratos.fstep6
            contratos.fstep5 = contratos.fstep6
            contratos.fstep4 = contratos.fstep6
            contratos.fstep3 = contratos.fstep6
            contratos.fstep2 = contratos.fstep6
            contratos.fstep1 = contratos.fstep6
            contratos.devuelto_por = contratos.current_user
            contratos.current_user = contratos.rcap
            contratos.status = "CAP"

            contratos.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")
    return HttpResponse("FAIL")


def numero_to_letras(numero):

    indicador = [
        ("", ""), ("MIL", "MIL"), ("MILLON", "MILLONES"), ("MIL", "MIL"), ("BILLON", "BILLONES")]

    entero = int(numero)

    decimal = int(round((numero - entero)*100))

    contador = 0

    numero_letras = ""

    while entero > 0:

        a = entero % 1000

        if contador == 0:

            en_letras = convierte_cifra(a, 1).strip()

        else:

            en_letras = convierte_cifra(a, 0).strip()

        if a == 0:

            numero_letras = en_letras+" "+numero_letras

        elif a == 1:

            if contador in (1, 3):

                numero_letras = indicador[contador][0]+" "+numero_letras

            else:

                numero_letras = en_letras+" " + \
                    indicador[contador][0]+" "+numero_letras

        else:

            numero_letras = en_letras+" " + \
                indicador[contador][1]+" "+numero_letras

        numero_letras = numero_letras.strip()

        contador = contador + 1

        entero = int(entero / 1000)

    numero_letras = numero_letras+" PESOS " + str(decimal) + "/100 M.N."
    return (numero_letras)


def convierte_cifra(numero, sw):

    lista_centana = ["", ("CIEN", "CIENTO"), "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
                     "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]

    lista_decena = ["", ("DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE", "DIECISEIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE"),
                    ("VEINTE", "VEINTI"), ("TREINTA",
                                           "TREINTA Y"), ("CUARENTA", "CUARENTA Y"),
                    ("CINCUENTA", "CINCUENTA Y"), ("SESENTA", "SESENTA Y"),
                    ("SETENTA", "SETENTA Y"), ("OCHENTA", "OCHENTA Y"),
                    ("NOVENTA", "NOVENTA Y")
                    ]

    lista_unidad = ["", ("UN", "UNO"), "DOS", "TRES",
                    "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]

    centena = int(numero / 100)

    decena = int((numero - (centena * 100))/10)

    unidad = int(numero - (centena * 100 + decena * 10))

    texto_centena = ""

    texto_decena = ""

    texto_unidad = ""

    texto_centena = lista_centana[centena]

    if centena == 1:

        if (decena + unidad) != 0:

            texto_centena = texto_centena[1]

        else:

            texto_centena = texto_centena[0]

    texto_decena = lista_decena[decena]

    if decena == 1:

        texto_decena = texto_decena[unidad]

    elif decena > 1:

        if unidad != 0:

            texto_decena = texto_decena[1]

        else:

            texto_decena = texto_decena[0]

    if decena != 1:

        texto_unidad = lista_unidad[unidad]

        if unidad == 1:

            texto_unidad = texto_unidad[sw]

    return "%s %s %s" % (texto_centena, texto_decena, texto_unidad)


def _add_field(run, field):
    """ add a field to a run
    """
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = field

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def _add_number_range(run):
    """ add a number range field to a run
    """
    _add_field(run, r'Page')


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def contratoGracont(request, id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.is_ajax and request.method == "POST":

        data = json.loads(request.body)
        # print(data)

        if not contratos.fcap:
            contratos.datecontrato = datetime.today()
            contratos.datecontrato_ini = data["enc_datecontrato_ini"]
            contratos.datecontrato_fin = data["enc_datecontrato_fin"]
            xiC = data["enc_importeContrato"]
            # print(xiC)
            # print(type(xiC))
            #contratos.importeContrato = data["enc_importeContrato"]

            contratos.npContrato = data["enc_npContrato"]
            #contratos.imppContrato = data["enc_imppContrato"]
            contratos.totalhorasContrato = data["enc_totalhorasContrato"]
            #contratos.testigoContrato1 = data["enc_testigoContrato1"]
            contratos.testigoContrato2 = data["enc_testigoContrato2"]

            contratos.save()
            return HttpResponse("OK")

        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def marcaContrato(request, id):
    tipocontratos = Tipocontrato.objects.filter(pk=id).first()

    if request.is_ajax and request.method == "POST":

        data = json.loads(request.body)
        # print(data)

        if Tipocontrato.tipoContrato:
            Tipocontrato.marcatipoContrato = True
            Tipocontrato.save()
            return HttpResponse("OK")

        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


def calcular_edad_anos(fecha_nacimiento):
    fecha_actual = date.today()
    resultado = fecha_actual.year - fecha_nacimiento.year
    resultado -= ((fecha_actual.month, fecha_actual.day) <
                  (fecha_nacimiento.month, fecha_nacimiento.day))
    return resultado


class DoctosDetDelete(SinPrivilegios, generic.DeleteView):
    permission_required = "cto.delete_doctos"
    model = Doctos
    template_name = "cto/doctos_det_del.html"
    context_object_name = 'obj'

    def get_success_url(self):
        contrato_id = self.kwargs['contrato_id']
        return reverse_lazy('cto:contrato_edit', kwargs={'contrato_id': contrato_id})


class PuestosView(SinPrivilegios, generic.ListView):
    model = Puestos
    template_name = "cto/puestos_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:puestos_list")
    permission_required = "cto.view_puestos"


class PuestosNew(VistaBaseCreate):
    model = Puestos
    template_name = "cto/puestos_form.html"
    form_class = PuestosForm
    success_url = reverse_lazy("cto:puestos_list")
    permission_required = "cto.add_puestos"


class PuestosEdit(VistaBaseEdit):
    model = Puestos
    template_name = "cto/puestos_form.html"
    form_class = PuestosForm
    success_url = reverse_lazy("cto:puestos_list")
    permission_required = "cto.change_puestos"
